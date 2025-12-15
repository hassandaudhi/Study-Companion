import uuid
import aiofiles
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from app.config.settings import settings
from fastapi import UploadFile


class FileService:
    """Service for handling file uploads and processing"""

    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.max_file_size = settings.MAX_FILE_SIZE

    async def save_file(self, file: UploadFile, user_id: int) -> dict:
        """
        Save uploaded file to storage
        
        Args:
            file: UploadFile object from FastAPI
            user_id: ID of the user uploading the file
        
        Returns:
            Dict with file info (filename, file_path, file_size, file_type)
        """
        try:
            # Validate file size
            file.file.seek(0, 2)  # Seek to end
            file_size = file.file.tell()
            file.file.seek(0)  # Reset to beginning

            if file_size > self.max_file_size:
                raise ValueError(f"File size exceeds maximum allowed size of {self.max_file_size} bytes")

            # Generate unique filename
            file_ext = Path(file.filename).suffix
            unique_filename = f"{uuid.uuid4()}{file_ext}"

            # Create user directory
            user_dir = self.upload_dir / str(user_id)
            user_dir.mkdir(parents=True, exist_ok=True)

            # Save file
            file_path = user_dir / unique_filename
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)

            return {
                "filename": unique_filename,
                "original_filename": file.filename,
                "file_path": str(file_path),
                "file_size": file_size,
                "file_type": file_ext.lstrip('.')
            }
        except Exception as e:
            raise Exception(f"Failed to save file: {str(e)}")

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Extracted text
            
        Raises:
            Exception: If PDF extraction fails
        """
        try:
            reader = PdfReader(file_path)
            if not reader.pages:
                raise ValueError("PDF file contains no pages")

            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            if not text.strip():
                raise ValueError("No text could be extracted from PDF")

            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Extracted text
            
        Raises:
            Exception: If DOCX extraction fails
        """
        try:
            doc = Document(file_path)
            if not doc.paragraphs:
                raise ValueError("DOCX file contains no paragraphs")

            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

            if not text.strip():
                raise ValueError("No text could be extracted from DOCX")

            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to TXT file
        
        Returns:
            Extracted text
            
        Raises:
            Exception: If TXT extraction fails
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read().strip()

            if not text:
                raise ValueError("TXT file is empty")

            return text
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read().strip()
                if not text:
                    raise ValueError("TXT file is empty")
                return text
            except Exception as e:
                raise Exception(f"Failed to extract text from TXT: {str(e)}")
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")

    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from file based on type
        
        Args:
            file_path: Path to file
            file_type: File extension (pdf, docx, txt)
        
        Returns:
            Extracted text
        """
        extractors = {
            'pdf': self.extract_text_from_pdf,
            'docx': self.extract_text_from_docx,
            'txt': self.extract_text_from_txt
        }

        extractor = extractors.get(file_type.lower())
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")

        return extractor(file_path)

    async def delete_file(self, file_path: str) -> bool:
        """
        Delete file from storage
        
        Args:
            file_path: Path to file
        
        Returns:
            True if successful
        """
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                return True
            return False
        except Exception as e:
            raise Exception(f"Failed to delete file: {str(e)}")

    def get_file_path(self, filename: str, user_id: int) -> str:
        """
        Get full file path
        
        Args:
            filename: Name of file
            user_id: User ID
        
        Returns:
            Full file path
        """
        return str(self.upload_dir / str(user_id) / filename)


# Singleton instance
file_service = FileService()
